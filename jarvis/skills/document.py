"""Document Generation Skill using fpdf and python-docx."""
import os
from fpdf import FPDF
from docx import Document
from jarvis.core.llm import LLM
from datetime import datetime

class DocumentGenerator:
    """Generates PDF and Word documents from AI content."""
    
    def __init__(self):
        self.llm = LLM()
        self.doc_dir = os.path.join(os.path.expanduser("~"), "Downloads", "JARVIS", "Documents")
        os.makedirs(self.doc_dir, exist_ok=True)

    def handle(self, query: str) -> str:
        """Route document generation request."""
        query_lower = query.lower()
        
        # Determine format
        file_format = "pdf"
        if "word" in query_lower or "docx" in query_lower:
            file_format = "docx"
            
        # Generate Content using Gemini
        try:
            prompt = f"""
            User Request: {query}
            
            Task: Generate the FULL content for this document. 
            Do NOT include "Here is the content" or markdown code blocks. 
            Just provide the raw text for the document title and body.
            Format:
            Title: [Title Here]
            [Body Content Here]
            """
            content = self.llm.chat(prompt)
            if not content:
                return "Failed to generate document content."
                
            # Parse Title and Body (Simple Heuristic)
            lines = content.split('\n')
            title = "Document"
            body = content
            
            if lines[0].strip().lower().startswith("title:"):
                title = lines[0].split(":", 1)[1].strip()
                body = "\n".join(lines[1:]).strip()
            
            # Generate File
            filename = f"{title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            
            if file_format == "pdf":
                filepath = self._create_pdf(title, body, filename)
            else:
                filepath = self._create_docx(title, body, filename)
                
            # Auto-open
            try:
                os.startfile(filepath)
                return f"Document created and opened: {filepath}"
            except Exception:
                return f"Document created at: {filepath}"

        except Exception as e:
            return f"Document generation failed: {e}"

    def _create_pdf(self, title: str, body: str, filename: str) -> str:
        """Create PDF file."""
        pdf = FPDF()
        pdf.add_page()
        
        # Title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, title, 0, 1, 'C')
        pdf.ln(10)
        
        # Body
        pdf.set_font("Arial", size=12)
        # simplistic handling of unicode - replace common incompatible chars
        safe_body = body.encode('latin-1', 'replace').decode('latin-1') 
        pdf.multi_cell(0, 10, safe_body)
        
        filepath = os.path.join(self.doc_dir, f"{filename}.pdf")
        pdf.output(filepath)
        return filepath

    def _create_docx(self, title: str, body: str, filename: str) -> str:
        """Create Word file."""
        doc = Document()
        doc.add_heading(title, 0)
        
        for paragraph in body.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
                
        filepath = os.path.join(self.doc_dir, f"{filename}.docx")
        doc.save(filepath)
        return filepath
